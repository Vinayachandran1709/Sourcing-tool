from fastapi import FastAPI, HTTPException, Depends, Request, Header, Response, UploadFile, File, Body
from fastapi.responses import StreamingResponse
import json
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

from typing import Annotated, List, Optional
from pydantic import BaseModel, ConfigDict
from datetime import datetime, timezone, timedelta
import logging
import sys
import asyncio
import math

# ⭐ CONFIGURE LOGGING (Windows-compatible, no emojis in console)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("talentbox.log", encoding='utf-8'),  # File supports UTF-8
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger(__name__)

# Configure StreamHandler to avoid emoji issues on Windows
for handler in logging.root.handlers:
    if isinstance(handler, logging.StreamHandler) and handler.stream == sys.stdout:
        handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

# Import routers
from auth_routes import router as auth_router
from candidate_auth_routes import router as candidate_auth_router
from routes.payment_routes import router as payment_router
from lists_routes import router as lists_router
from email_settings_routes import router as email_settings_router
from feedback_routes import router as feedback_router

# Import services
from filter_service import FilterService
from usage_service import UsageService
from database import get_db
from models import User, Profile, EmailOutreach
from auth_middleware import get_current_user
from profile_cache_service import ProfileCacheService
from email_service import EmailService
from redis_service import init_redis, get_redis_stats, get_cached_search, set_cached_search, hash_filters, get_cached_jd_parse, set_cached_jd_parse
from ai_service import generate_profile_summary, parse_job_description, check_groq_status
import openai
from config import GITHUB_TOKEN

openai.api_key = os.getenv("OPENAI_API_KEY")

# ===== ANNOTATED DEPENDENCY TYPES =====
CurrentUser = Annotated[User, Depends(get_current_user)]
DbSession = Annotated[Session, Depends(get_db)]

# ===== INITIALIZE FASTAPI APP =====

# ===== CONFIGURE CORS ORIGINS FIRST =====

_raw_origins = os.getenv("CORS_ORIGINS", "http://localhost:3000").split(",")
CORS_ORIGINS = list({o.strip() for o in _raw_origins if o.strip()} | {
    "https://www.talentbox.co",
    "https://talentbox.co",
    "http://localhost:3000",
})


# ===== INITIALIZE FASTAPI APP =====

limiter = Limiter(key_func=get_remote_address)

app = FastAPI(
    title="TalentBox API",
    version="2.0.1",  # ✅ Updated version for Resend integration
    description="API for GitHub developer sourcing and recruitment"
)

SUSPICIOUS_PATH_FRAGMENTS = (
    ".env", ".git", ".aws", "config.json", "config.env",
    "settings.json", "buildmanifest", "manifest.json",
    "wp-admin", "phpinfo", "/admin/.", "graphql",
    ".env.local", ".env.production", ".env.development",
    ".env.backup", ".env.bak", ".env.old", ".env.example",
)

@app.middleware("http")
async def block_scanner_probes(request: Request, call_next):
    path_lower = request.url.path.lower()
    if any(frag in path_lower for frag in SUSPICIOUS_PATH_FRAGMENTS):
        return Response(status_code=404)
    return await call_next(request)

# Log initialization
logger.info("TalentBox API initialized - Version 2.0.1 (Resend Integration)")
logger.info(f"Environment: {os.getenv('ENVIRONMENT', 'development')}")
logger.info(f"CORS origins: {CORS_ORIGINS}")

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

@app.on_event("startup")
async def startup_event():
    """Validate critical configurations on startup"""
    logger.info("TalentBox API starting up...")
    
    # Validate GitHub Token
    from github_service import GITHUB_TOKEN
    if not GITHUB_TOKEN:
        logger.error("CRITICAL: GITHUB_TOKEN not found in .env!")
        logger.error("   GitHub API searches will fail")
        logger.error("   Only database cache will work")
    else:
        logger.info("GitHub token found")
    
    # Validate Resend API Key
    resend_key = os.getenv("RESEND_API_KEY")
    if not resend_key:
        logger.warning("RESEND_API_KEY not found - email features will not work")
    else:
        logger.info("Resend API key configured")

    logger.info("Skipping startup DB health check (lazy connect on first request)")

    # Auto-run payment migration (safe - uses IF NOT EXISTS)
    if os.getenv("RUN_MIGRATIONS_ON_STARTUP", "false").lower() == "true":
        try:
            from sqlalchemy import text as sa_text
            from database import engine as db_engine
            with db_engine.connect() as conn:
                # Add payment columns to users table
                conn.execute(sa_text("""
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS razorpay_customer_id VARCHAR(50);
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS razorpay_order_id VARCHAR(50);
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS auto_renew BOOLEAN DEFAULT TRUE;
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS payment_method VARCHAR(20);
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS last_payment_date TIMESTAMP WITH TIME ZONE;
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS subscription_amount DECIMAL(10, 2) DEFAULT 0;
                    ALTER TABLE users ADD COLUMN IF NOT EXISTS currency VARCHAR(3) DEFAULT 'USD';
                """))

                # Create payment_history table
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS payment_history (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                        razorpay_order_id VARCHAR(50) NOT NULL,
                        razorpay_payment_id VARCHAR(50),
                        razorpay_signature VARCHAR(255),
                        amount DECIMAL(10, 2) NOT NULL,
                        currency VARCHAR(3) DEFAULT 'USD',
                        amount_inr DECIMAL(10, 2),
                        plan_name VARCHAR(50) NOT NULL,
                        billing_cycle VARCHAR(20) NOT NULL,
                        status VARCHAR(20) DEFAULT 'created',
                        payment_method VARCHAR(50),
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                        paid_at TIMESTAMP WITH TIME ZONE,
                        receipt VARCHAR(100),
                        notes JSONB,
                        error_message TEXT,
                        CONSTRAINT unique_razorpay_order UNIQUE (razorpay_order_id)
                    );
                    CREATE INDEX IF NOT EXISTS idx_payment_history_user_id ON payment_history(user_id);
                    CREATE INDEX IF NOT EXISTS idx_payment_history_status ON payment_history(status);
                """))

                # Create subscription_events table
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS subscription_events (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                        event_type VARCHAR(50) NOT NULL,
                        old_plan VARCHAR(50),
                        new_plan VARCHAR(50),
                        old_status VARCHAR(20),
                        new_status VARCHAR(20),
                        triggered_by VARCHAR(50),
                        metadata JSONB,
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                    );
                    CREATE INDEX IF NOT EXISTS idx_subscription_events_user_id ON subscription_events(user_id);
                """))

                # Add GraphQL + rate limiting + smart refresh columns and tables
                conn.execute(sa_text("""
                    ALTER TABLE profiles ADD COLUMN IF NOT EXISTS followers INTEGER DEFAULT 0;
                    ALTER TABLE profiles ADD COLUMN IF NOT EXISTS is_hireable BOOLEAN DEFAULT FALSE;
                    ALTER TABLE profiles ADD COLUMN IF NOT EXISTS refresh_category VARCHAR(20) DEFAULT 'dormant';
                    ALTER TABLE profiles ADD COLUMN IF NOT EXISTS last_refreshed_at TIMESTAMP WITH TIME ZONE;
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS search_locks (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                        locked_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                        expires_at TIMESTAMP WITH TIME ZONE NOT NULL,
                        search_completed BOOLEAN DEFAULT FALSE,
                        CONSTRAINT unique_user_lock UNIQUE (user_id)
                    );
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS rate_limit_events (
                        id SERIAL PRIMARY KEY,
                        event_type VARCHAR(20) NOT NULL,
                        status_code INTEGER NOT NULL,
                        retry_after INTEGER,
                        rate_limit_remaining INTEGER,
                        rate_limit_resource VARCHAR(50),
                        endpoint VARCHAR(200),
                        occurred_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                    );
                    CREATE INDEX IF NOT EXISTS idx_rate_limit_events_time ON rate_limit_events(occurred_at);
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS profile_unlocks (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                        profile_id INTEGER NOT NULL REFERENCES profiles(id) ON DELETE CASCADE,
                        unlocked_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                    );
                    CREATE INDEX IF NOT EXISTS idx_profile_unlocks_profile ON profile_unlocks(profile_id);
                    CREATE INDEX IF NOT EXISTS idx_profile_unlocks_time ON profile_unlocks(unlocked_at);
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS refresh_job_log (
                        id SERIAL PRIMARY KEY,
                        started_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                        completed_at TIMESTAMP WITH TIME ZONE,
                        profiles_refreshed INTEGER DEFAULT 0,
                        profiles_failed INTEGER DEFAULT 0,
                        status VARCHAR(20) DEFAULT 'running'
                    );
                """))

                # Backfill refresh categories for existing profiles
                conn.execute(sa_text("""
                    UPDATE profiles SET refresh_category = 'active' WHERE contributions_last_year >= 300 AND refresh_category = 'dormant';
                    UPDATE profiles SET refresh_category = 'moderate' WHERE contributions_last_year >= 100 AND contributions_last_year < 300 AND refresh_category = 'dormant';
                """))

                # Add account_type and estimated_experience_years to github_developers
                conn.execute(sa_text("""
                    ALTER TABLE github_developers ADD COLUMN IF NOT EXISTS account_type VARCHAR(20) DEFAULT 'User';
                    ALTER TABLE github_developers ADD COLUMN IF NOT EXISTS estimated_experience_years INTEGER DEFAULT 0;
                    CREATE INDEX IF NOT EXISTS idx_github_developers_account_type ON github_developers(account_type);
                """))

                # Backfill known org accounts using heuristics
                conn.execute(sa_text("""
                    UPDATE github_developers SET account_type = 'Organization'
                    WHERE (account_type IS NULL OR account_type = 'User')
                    AND (
                        bio ILIKE '%organization%'
                        OR bio ILIKE '%open source org%'
                        OR name ILIKE '%Inc%'
                        OR name ILIKE '%Corp%'
                        OR name ILIKE '%Labs%'
                        OR name ILIKE '%Foundation%'
                        OR name ILIKE '%Technologies%'
                        OR (followers > 5000 AND public_repos > 100 AND email IS NULL)
                    );
                """))

                # Backfill estimated_experience_years from github_created_at
                conn.execute(sa_text("""
                    UPDATE github_developers
                    SET estimated_experience_years = EXTRACT(YEAR FROM NOW()) - EXTRACT(YEAR FROM github_created_at)
                    WHERE github_created_at IS NOT NULL AND (estimated_experience_years IS NULL OR estimated_experience_years = 0);
                """))

                # DevCard profiles table
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS devcard_profiles (
                        id SERIAL PRIMARY KEY,
                        github_username VARCHAR(255) UNIQUE NOT NULL,
                        display_name VARCHAR(255),
                        avatar_url VARCHAR(500),
                        bio TEXT,
                        location VARCHAR(255),
                        detected_role VARCHAR(100),
                        seniority_level VARCHAR(50),
                        primary_languages TEXT[],
                        language_percentages JSONB DEFAULT '{}',
                        top_projects JSONB DEFAULT '[]',
                        contribution_stats JSONB DEFAULT '{}',
                        ai_summary TEXT,
                        experience_history JSONB DEFAULT '[]',
                        developer_score INTEGER DEFAULT 0,
                        estimated_experience_years INTEGER DEFAULT 0,
                        resume_uploaded BOOLEAN DEFAULT FALSE,
                        email VARCHAR(255),
                        linkedin_url VARCHAR(500),
                        phone VARCHAR(50),
                        views_count INTEGER DEFAULT 0,
                        is_published BOOLEAN DEFAULT TRUE,
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS idx_devcard_username ON devcard_profiles(github_username);
                """))

                # Candidate waitlist (lead capture from DevCard creation)
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS candidate_waitlist (
                        id SERIAL PRIMARY KEY,
                        email VARCHAR(255),
                        name VARCHAR(255),
                        github_username VARCHAR(255),
                        linkedin_url VARCHAR(500),
                        phone VARCHAR(50),
                        source VARCHAR(50) DEFAULT 'devcard',
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS idx_waitlist_email ON candidate_waitlist(email);
                """))

                # Free hire requests (companies posting jobs for free matching)
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS free_hire_requests (
                        id SERIAL PRIMARY KEY,
                        company_email VARCHAR(255) NOT NULL,
                        company_name VARCHAR(255),
                        contact_name VARCHAR(255),
                        job_title VARCHAR(255) NOT NULL,
                        job_description TEXT,
                        required_skills TEXT[],
                        preferred_location VARCHAR(255),
                        experience_min INTEGER DEFAULT 0,
                        remote_ok BOOLEAN DEFAULT FALSE,
                        matched_profiles_count INTEGER DEFAULT 0,
                        email_sent BOOLEAN DEFAULT FALSE,
                        email_sent_at TIMESTAMP WITH TIME ZONE,
                        status VARCHAR(50) DEFAULT 'pending',
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS idx_hire_requests_email ON free_hire_requests(company_email);
                    CREATE INDEX IF NOT EXISTS idx_hire_requests_status ON free_hire_requests(status);
                """))

                # Extend free_hire_requests with parsed JD detail columns
                conn.execute(sa_text("""
                    ALTER TABLE free_hire_requests
                        ADD COLUMN IF NOT EXISTS jd_source VARCHAR(20) DEFAULT 'pasted',
                        ADD COLUMN IF NOT EXISTS jd_filename VARCHAR(255),
                        ADD COLUMN IF NOT EXISTS jd_parsed_role VARCHAR(255),
                        ADD COLUMN IF NOT EXISTS jd_parsed_skills TEXT[],
                        ADD COLUMN IF NOT EXISTS jd_parsed_location VARCHAR(255),
                        ADD COLUMN IF NOT EXISTS jd_parsed_experience INTEGER,
                        ADD COLUMN IF NOT EXISTS jd_word_count INTEGER,
                        ADD COLUMN IF NOT EXISTS jd_char_count INTEGER;
                """))

                # PIVOT TABLES
                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS candidates (
                        id SERIAL PRIMARY KEY,
                        email VARCHAR(255) UNIQUE NOT NULL,
                        password_hash VARCHAR(255) NOT NULL,
                        name VARCHAR(255) NOT NULL,
                        phone VARCHAR(50),
                        github_username VARCHAR(255),
                        linkedin_url VARCHAR(500),
                        portfolio_url VARCHAR(500),
                        resume_path VARCHAR(500),
                        avatar_url VARCHAR(500),
                        location VARCHAR(255),
                        onboarding_status VARCHAR(50) DEFAULT 'pending',
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS ix_candidates_email ON candidates(email);
                    CREATE INDEX IF NOT EXISTS ix_candidates_github_username ON candidates(github_username);
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS candidate_profiles (
                        id SERIAL PRIMARY KEY,
                        candidate_id INTEGER UNIQUE NOT NULL REFERENCES candidates(id) ON DELETE CASCADE,
                        github_analysis JSONB DEFAULT '{}',
                        resume_data JSONB DEFAULT '{}',
                        portfolio_data JSONB DEFAULT '{}',
                        conversation_data JSONB DEFAULT '{}',
                        career_preferences JSONB DEFAULT '{}',
                        technical_assessment JSONB DEFAULT '{}',
                        ai_summary TEXT,
                        detected_role VARCHAR(100),
                        detected_roles JSONB DEFAULT '[]',
                        seniority_level VARCHAR(50),
                        skills JSONB DEFAULT '[]',
                        engineering_maturity_score INTEGER DEFAULT 0,
                        profile_quality_score INTEGER DEFAULT 0,
                        match_ready BOOLEAN DEFAULT FALSE,
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS discovered_startups (
                        id SERIAL PRIMARY KEY,
                        company_name VARCHAR(255) NOT NULL,
                        domain VARCHAR(255) UNIQUE,
                        description TEXT,
                        logo_url VARCHAR(500),
                        funding_stage VARCHAR(50),
                        funding_amount VARCHAR(100),
                        investors JSONB DEFAULT '[]',
                        team_size VARCHAR(50),
                        category VARCHAR(100),
                        location VARCHAR(255),
                        remote_policy VARCHAR(50),
                        ats_provider VARCHAR(50),
                        ats_slug VARCHAR(255),
                        careers_url VARCHAR(500),
                        source VARCHAR(50) DEFAULT 'manual',
                        is_active BOOLEAN DEFAULT TRUE,
                        last_crawled_at TIMESTAMP WITH TIME ZONE,
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS ix_discovered_startups_domain ON discovered_startups(domain);
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS job_postings (
                        id SERIAL PRIMARY KEY,
                        external_id VARCHAR(255),
                        startup_id INTEGER REFERENCES discovered_startups(id) ON DELETE SET NULL,
                        company_name VARCHAR(255) NOT NULL,
                        company_domain VARCHAR(255),
                        company_logo_url VARCHAR(500),
                        funding_stage VARCHAR(50),
                        investors_summary VARCHAR(500),
                        title VARCHAR(255) NOT NULL,
                        department VARCHAR(100),
                        description_text TEXT,
                        description_html TEXT,
                        location VARCHAR(255),
                        remote_policy VARCHAR(50),
                        seniority_level VARCHAR(50),
                        required_stack JSONB DEFAULT '[]',
                        salary_min INTEGER,
                        salary_max INTEGER,
                        salary_currency VARCHAR(10),
                        ats_provider VARCHAR(50),
                        ats_url VARCHAR(500),
                        apply_url VARCHAR(500),
                        source VARCHAR(50) DEFAULT 'manual',
                        quality_score INTEGER DEFAULT 50,
                        is_engineering BOOLEAN DEFAULT TRUE,
                        is_active BOOLEAN DEFAULT TRUE,
                        is_visible_on_homepage BOOLEAN DEFAULT TRUE,
                        posted_by_email VARCHAR(255),
                        posted_by_name VARCHAR(255),
                        posted_by_phone VARCHAR(50),
                        company_follow_up_data JSONB DEFAULT '{}',
                        first_seen_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        last_seen_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                    CREATE INDEX IF NOT EXISTS ix_job_postings_active_visible ON job_postings(is_active, is_visible_on_homepage);
                    CREATE INDEX IF NOT EXISTS ix_job_postings_company_title ON job_postings(company_domain, title);
                """))

                conn.execute(sa_text("""
                    CREATE TABLE IF NOT EXISTS job_matches (
                        id SERIAL PRIMARY KEY,
                        job_posting_id INTEGER NOT NULL REFERENCES job_postings(id) ON DELETE CASCADE,
                        candidate_id INTEGER REFERENCES candidates(id) ON DELETE CASCADE,
                        github_developer_id INTEGER REFERENCES github_developers(id) ON DELETE SET NULL,
                        match_score INTEGER DEFAULT 0,
                        match_reasons JSONB DEFAULT '{}',
                        status VARCHAR(50) DEFAULT 'matched',
                        sent_at TIMESTAMP WITH TIME ZONE,
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    );
                """))

                conn.commit()
            logger.info("Payment tables verified/created")
        except Exception as e:
            logger.error(f"Payment migration error: {e}")
    else:
        logger.info("Skipping startup migrations (set RUN_MIGRATIONS_ON_STARTUP=true to apply schema changes)")

    # Validate JWT Secret
    from auth_middleware import SECRET_KEY
    if not SECRET_KEY:
        logger.error("CRITICAL: JWT_SECRET_KEY not found!")
    else:
        logger.info("JWT secret configured")
    
    # Initialize background scheduler for nightly profile refresh
    try:
        from scheduler import init_scheduler
        init_scheduler()
        logger.info("Background scheduler initialized")
    except Exception as e:
        logger.error(f"Scheduler init failed: {e}")

    # Initialize Redis cache
    await init_redis()

    logger.info("Startup validation complete\n")


@app.on_event("shutdown")
async def shutdown_event():
    """Gracefully shut down scheduler on app shutdown."""
    try:
        from scheduler import shutdown_scheduler
        shutdown_scheduler()
    except Exception:
        pass


# ⭐ GLOBAL ERROR HANDLERS

from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Handle validation errors with user-friendly messages"""
    logger.warning(f"Validation error on {request.url.path}: {exc.errors()}")
    
    # Convert errors to JSON-serializable format
    errors = []
    for error in exc.errors():
        error_dict = {
            "type": error.get("type"),
            "loc": error.get("loc"),
            "msg": error.get("msg"),
            "input": str(error.get("input")) if error.get("input") is not None else None
        }
        errors.append(error_dict)
    
    return JSONResponse(
        status_code=422,
        content={
            "success": False,
            "error": "VALIDATION_ERROR",
            "message": "Invalid input data",
            "details": errors
        }
    )

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """Handle HTTP exceptions with consistent format"""
    logger.warning(f"HTTP {exc.status_code} on {request.url.path}: {exc.detail}")
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "success": False,
            "error": exc.detail if isinstance(exc.detail, str) else exc.detail.get("error", "ERROR"),
            "message": exc.detail if isinstance(exc.detail, str) else exc.detail.get("message", str(exc.detail))
        }
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Handle unexpected errors without exposing internals"""
    logger.error(f"Unexpected error on {request.url.path}: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "success": False,
            "error": "INTERNAL_SERVER_ERROR",
            "message": "An unexpected error occurred. Please try again later."
        }
    )


# ===== ADMIN ENDPOINTS =====
#just for reference
@app.get("/admin/cache-stats")
async def cache_stats(x_admin_key: str = Header(..., alias="X-Admin-Key")):
    """Return Redis cache statistics. Requires X-Admin-Key header."""
    admin_secret = os.getenv("ADMIN_SECRET_KEY")
    if not admin_secret or x_admin_key != admin_secret:
        raise HTTPException(status_code=403, detail="Forbidden")
    return await get_redis_stats()


# ===== CORS MIDDLEWARE =====

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===== INCLUDE ROUTERS =====

app.include_router(auth_router)
app.include_router(candidate_auth_router)
app.include_router(payment_router)
app.include_router(lists_router)
app.include_router(email_settings_router)
app.include_router(feedback_router)

# ===== REQUEST/RESPONSE MODELS =====

class SearchRequest(BaseModel):
    """Simplified search request - role and location based"""
    role: Optional[str] = None
    location: Optional[str] = None
    languages: Optional[List[str]] = None
    min_score: Optional[int] = 0
    page: Optional[int] = 1
    per_page: Optional[int] = 20
    has_email: Optional[bool] = None
    min_experience: Optional[int] = None

class ProfileResponse(BaseModel):
    """Response model for profile data"""
    model_config = ConfigDict(from_attributes=True)
    
    id: int
    github_username: str
    name: Optional[str]
    email: Optional[str]
    location: Optional[str]
    bio: Optional[str]
    public_repos: int
    primary_language: Optional[str]
    contributions_last_year: int
    portfolio_url: Optional[str]
    avatar_url: Optional[str]
    selected: bool
    total_stars: int
    developer_score: int
    languages_data: Optional[dict] = None
    top_repos: Optional[list] = None
    last_active_date: Optional[datetime] = None
    last_fetched: Optional[datetime] = None


class EmailRequest(BaseModel):
    """Request model for sending emails"""
    profile_ids: List[int]


# ===== ROOT ENDPOINT =====

@app.get("/")
def root():
    """Welcome endpoint"""
    return {
        "message": "TalentBox API - Developer Sourcing Platform",
        "version": "2.0.1",
        "status": "operational",
        "docs": "/docs"
    }


# ===== SEARCH ENDPOINT =====

@app.post("/api/search-profiles")
async def search_profiles(
    request: Request,
    search: SearchRequest,
    current_user: CurrentUser,
    db: DbSession,
):
    """Search for developer profiles — streams results progressively via SSE."""
    user_id = current_user.id

    logger.info(f"Search request from user {user_id}: role={search.role}, location={search.location}")

    # Check usage limits
    try:
        UsageService.check_limit(db, user_id, "search")
    except HTTPException as e:
        return {"error": e.detail, "limit_reached": True}

    role = search.role
    location = search.location
    languages = search.languages
    min_score = search.min_score or 0
    is_paid = current_user.subscription_status == "active"

    def _first_name(full_name):
        if full_name:
            return full_name.split()[0]
        return None

    def _profile_obj_to_dict(profile):
        return {
            "id": profile.id,
            "github_username": profile.github_username,
            "name": _first_name(profile.name),
            "email": profile.email,
            "location": profile.location,
            "bio": profile.bio,
            "public_repos": profile.public_repos,
            "primary_language": profile.primary_language,
            "total_stars": getattr(profile, 'total_stars', 0),
            "developer_score": getattr(profile, 'developer_score', 0),
            "avatar_url": getattr(profile, 'avatar_url', None),
            "total_contributions": getattr(profile, 'contributions_last_year', 0),
            "followers": getattr(profile, 'followers', 0),
            "languages_data": getattr(profile, 'languages_data', None),
            "top_repos": getattr(profile, 'top_repos', None),
            "selected": False,
        }

    async def event_stream():
        from archive_search_service import search_developers, search_developers_batched, count_developers, developer_to_dict
        from github_integration_service import GitHubIntegrationService

        seen_usernames = set()
        total_sent = 0

        # --- STEP 1: Archive search (fast DB query, batched) ---
        try:
            total_matching = count_developers(db, role=role, location=location, languages=languages, min_score=min_score)
            yield f"data: {json.dumps({'type': 'count', 'total_matching': total_matching})}\n\n"

            batch_index = 0
            for batch in search_developers_batched(
                db=db, role=role, location=location, languages=languages,
                batch_size=500, min_score=min_score,
            ):
                batch_dicts = [developer_to_dict(dev) for dev in batch]
                for p in batch_dicts:
                    p["name"] = _first_name(p.get("name"))
                    seen_usernames.add(p["github_username"])

                total_sent += len(batch_dicts)

                if batch_index == 0:
                    yield f"data: {json.dumps({'type': 'profiles', 'profiles': batch_dicts})}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'new_profiles', 'profiles': batch_dicts})}\n\n"

                yield f"data: {json.dumps({'type': 'progress', 'loaded': total_sent, 'total': total_matching})}\n\n"
                batch_index += 1

            logger.info(f"Streamed {total_sent} archive profiles in {batch_index} batches")
        except Exception as e:
            logger.error(f"Archive search failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'profiles', 'profiles': []})}\n\n"

        # --- STEP 2: Supplement from Profile table if archive was small ---
        if total_sent < 500:
            yield f"data: {json.dumps({'type': 'status', 'message': 'Fetching more profiles...'})}\n\n"
            try:
                filters = {"role": role, "location": location, "min_score": min_score}
                cached_profiles = GitHubIntegrationService._search_database(db, filters)
                new_dicts = []
                for profile in cached_profiles:
                    if profile.github_username not in seen_usernames:
                        seen_usernames.add(profile.github_username)
                        new_dicts.append(_profile_obj_to_dict(profile))

                if new_dicts:
                    total_sent += len(new_dicts)
                    yield f"data: {json.dumps({'type': 'new_profiles', 'profiles': new_dicts})}\n\n"
                    logger.info(f"Streamed {len(new_dicts)} supplement profiles")
            except Exception as e:
                logger.error(f"Profile table supplement failed: {e}", exc_info=True)

        # --- STEP 3: Live GitHub API supplement for paid users ---
        if total_sent < 200 and is_paid:
            yield f"data: {json.dumps({'type': 'status', 'message': 'Searching GitHub for more...'})}\n\n"
            try:
                from github_integration_service import check_github_rate_limit
                await check_github_rate_limit()
                new_profiles = await GitHubIntegrationService._fetch_from_github(
                    db, None, location, 0, 250, 200 - total_sent
                )
                api_dicts = []
                for profile in new_profiles:
                    if profile.github_username not in seen_usernames:
                        seen_usernames.add(profile.github_username)
                        api_dicts.append(_profile_obj_to_dict(profile))
                if api_dicts:
                    total_sent += len(api_dicts)
                    yield f"data: {json.dumps({'type': 'new_profiles', 'profiles': api_dicts})}\n\n"
                    logger.info(f"Streamed {len(api_dicts)} GitHub API profiles")
            except Exception as e:
                logger.error(f"GitHub API supplement failed: {e}", exc_info=True)

        # --- DONE ---
        UsageService.log_usage(db, user_id, "search", {"role": role, "location": location})
        yield f"data: {json.dumps({'type': 'complete', 'total': total_sent})}\n\n"
        logger.info(f"Search complete: {total_sent} total profiles")

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ===== CANDIDATE ENDPOINTS =====

from candidate_auth_middleware import get_current_candidate
from models import CandidateProfile

@app.post("/api/candidate/import-data")
async def trigger_candidate_import(
    db: DbSession,
    current_candidate = Depends(get_current_candidate),
):
    """Trigger the background data import pipeline."""
    from data_import_service import run_candidate_import
    result = await run_candidate_import(db, current_candidate.id)
    return result

@app.get("/api/candidate/import-status")
async def get_candidate_import_status(
    db: DbSession,
    current_candidate = Depends(get_current_candidate),
):
    """Get the current onboarding status."""
    return {
        "success": True,
        "onboarding_status": current_candidate.onboarding_status
    }

@app.get("/api/candidate/profile")
async def get_candidate_profile(
    db: DbSession,
    current_candidate = Depends(get_current_candidate),
):
    """Get current candidate profile data \u2014 full response for dashboard."""
    profile = db.query(CandidateProfile).filter(
        CandidateProfile.candidate_id == current_candidate.id
    ).first()
    
    return {
        "success": True,
        "candidate": {
            "id": current_candidate.id,
            "name": current_candidate.name,
            "email": current_candidate.email,
            "github_username": current_candidate.github_username,
            "phone": current_candidate.phone,
            "location": current_candidate.location,
            "avatar_url": current_candidate.avatar_url,
            "linkedin_url": current_candidate.linkedin_url,
            "portfolio_url": current_candidate.portfolio_url,
            "onboarding_status": current_candidate.onboarding_status,
            "profile": {
                "github_analysis": profile.github_analysis if profile else {},
                "resume_data": profile.resume_data if profile else {},
                "portfolio_data": profile.portfolio_data if profile else {},
                "conversation_data": profile.conversation_data if profile else {},
                "career_preferences": profile.career_preferences if profile else {},
                "technical_assessment": profile.technical_assessment if profile else {},
                "ai_summary": profile.ai_summary if profile else None,
                "detected_role": profile.detected_role if profile else None,
                "detected_roles": profile.detected_roles if profile else [],
                "seniority_level": profile.seniority_level if profile else None,
                "skills": profile.skills if profile else [],
                "engineering_maturity_score": profile.engineering_maturity_score if profile else 0,
                "profile_quality_score": profile.profile_quality_score if profile else 0,
                "match_ready": profile.match_ready if profile else False,
            }
        }
    }

@app.post("/api/candidate/upload-resume")
async def upload_candidate_resume(
    db: DbSession,
    file: UploadFile = File(...),
    current_candidate = Depends(get_current_candidate),
):
    """Upload resume and parse basic text data."""
    import io
    try:
        filename = file.filename or ""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext not in ("pdf", "docx", "doc", "txt"):
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT.")

        raw = await file.read()
        if len(raw) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large (max 5 MB).")

        text = ""
        if ext == "txt":
            text = raw.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(raw))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as pdf_err:
                logger.error(f"PDF extraction error: {pdf_err}")
                raise HTTPException(status_code=422, detail="Could not extract text from PDF.")
        elif ext in ("docx", "doc"):
            try:
                import docx as _docx
                doc = _docx.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as docx_err:
                logger.error(f"DOCX extraction error: {docx_err}")
                raise HTTPException(status_code=422, detail="Could not extract text from DOCX.")

        text = text.strip()
        if len(text) < 30:
            raise HTTPException(status_code=422, detail="File appears to be empty or unreadable.")
            
        # Update candidate resume path
        current_candidate.resume_path = filename
        
        # Update profile with resume data
        profile = db.query(CandidateProfile).filter(CandidateProfile.candidate_id == current_candidate.id).first()
        if not profile:
            profile = CandidateProfile(candidate_id=current_candidate.id)
            db.add(profile)
            
        resume_data = {
            "raw_text": text[:50000],  # Limit size
            "filename": filename,
            "word_count": len(text.split())
        }
        
        # Merge if exists
        existing_data = profile.resume_data or {}
        existing_data.update(resume_data)
        profile.resume_data = existing_data
        
        db.commit()

        return {
            "success": True,
            "filename": filename,
            "word_count": resume_data["word_count"],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume upload error: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload and parse resume.")

@app.post("/api/candidate/conversation/start")
async def start_candidate_conversation(
    db: DbSession,
    current_candidate = Depends(get_current_candidate),
):
    """Start the AI conversation for a candidate."""
    from candidate_conversation_agent import start_conversation
    
    profile = db.query(CandidateProfile).filter(
        CandidateProfile.candidate_id == current_candidate.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found. Complete data import first.")
    
    # Check if conversation already exists and is complete
    existing_conv = profile.conversation_data or {}
    if existing_conv.get("completed_at"):
        raise HTTPException(status_code=400, detail="Conversation already completed.")
    
    # Check if candidate has at least GitHub analysis
    if not profile.github_analysis:
        raise HTTPException(status_code=400, detail="Please complete the GitHub profile import first.")
    
    result = start_conversation(profile.github_analysis, profile.resume_data or {})
    
    profile.conversation_data = result["conversation_data"]
    current_candidate.onboarding_status = "conversation_started"
    db.commit()
    
    return {
        "success": True,
        "message": result["ai_message"],
        "state": result["conversation_data"]["state"],
        "stages_completed": result["conversation_data"]["stages_completed"]
    }

@app.post("/api/candidate/conversation/message")
async def send_candidate_message(
    db: DbSession,
    body: dict = Body(...),
    current_candidate = Depends(get_current_candidate),
):
    """Process a candidate's message in the conversation."""
    from candidate_conversation_agent import process_candidate_message
    
    candidate_message = body.get("message", "").strip()
    if not candidate_message:
        raise HTTPException(status_code=400, detail="Message cannot be empty.")
    if len(candidate_message) > 5000:
        raise HTTPException(status_code=400, detail="Message too long (max 5000 characters).")
    
    profile = db.query(CandidateProfile).filter(
        CandidateProfile.candidate_id == current_candidate.id
    ).first()
    
    if not profile or not profile.conversation_data:
        raise HTTPException(status_code=400, detail="No active conversation. Start one first.")
    
    conv_data = profile.conversation_data
    if conv_data.get("completed_at"):
        return {
            "success": True,
            "message": "Your conversation is already complete. Check your dashboard.",
            "is_complete": True,
            "state": "summary"
        }
    
    # Get last AI message timestamp for response timing
    last_ai_msg = None
    for m in reversed(conv_data.get("messages", [])):
        if m["role"] == "assistant":
            last_ai_msg = m.get("timestamp")
            break
    
    updated_conv, ai_response, is_complete = await process_candidate_message(
        candidate_message=candidate_message,
        conversation_data=conv_data,
        github_analysis=profile.github_analysis or {},
        resume_data=profile.resume_data or {},
        last_ai_timestamp=last_ai_msg
    )
    
    # Reassign the dictionary back to the SQLAlchemy object to trigger JSONB mutation tracking
    from sqlalchemy.orm.attributes import flag_modified
    profile.conversation_data = updated_conv
    flag_modified(profile, "conversation_data")
    
    if is_complete:
        current_candidate.onboarding_status = "conversation_complete"
    
    db.commit()
    
    return {
        "success": True,
        "message": ai_response,
        "is_complete": is_complete,
        "state": updated_conv.get("state"),
        "technical_questions_asked": updated_conv.get("technical_questions_asked", 0),
        "career_questions_asked": updated_conv.get("career_questions_asked", 0),
        "stages_completed": updated_conv.get("stages_completed", [])
    }

@app.post("/api/candidate/conversation/complete")
async def complete_candidate_conversation(
    db: DbSession,
    current_candidate = Depends(get_current_candidate),
):
    """Finalize the conversation and extract structured profile data."""
    from candidate_conversation_agent import extract_final_profile
    
    profile = db.query(CandidateProfile).filter(
        CandidateProfile.candidate_id == current_candidate.id
    ).first()
    
    if not profile or not profile.conversation_data:
        raise HTTPException(status_code=400, detail="No conversation data found.")
    
    conv_data = profile.conversation_data
    if not conv_data.get("completed_at"):
        raise HTTPException(status_code=400, detail="Conversation not yet complete.")
    
    # Extract final profile
    extracted = await extract_final_profile(
        conv_data,
        profile.github_analysis or {},
        profile.resume_data or {}
    )
    
    profile.career_preferences = extracted.get("career_preferences", {})
    profile.technical_assessment = extracted.get("technical_assessment", {})
    
    if extracted.get("updated_summary"):
        profile.ai_summary = extracted["updated_summary"]
    
    # Update seniority if the interview assessment differs from GitHub-only assessment
    interview_seniority = extracted.get("technical_assessment", {}).get("seniority_assessment")
    if interview_seniority:
        profile.seniority_level = interview_seniority
    
    # Mark as match-ready
    profile.match_ready = True
    current_candidate.onboarding_status = "profile_ready"
    
    # Recalculate profile quality score
    score = profile.profile_quality_score or 0
    score = min(score + 20, 100)  # Conversation adds 20 points
    profile.profile_quality_score = score
    
    db.commit()
    
    return {
        "success": True,
        "message": "Profile finalized successfully!",
        "career_preferences": profile.career_preferences,
        "technical_assessment": profile.technical_assessment,
        "match_ready": True
    }


# ===== PAGINATED SEARCH ENDPOINT =====

@app.post("/api/search")
@limiter.limit("30/minute")
async def search_paginated(
    request: Request,
    search: SearchRequest,
    current_user: CurrentUser,
    db: DbSession,
):
    """Paginated search for developer profiles — returns instant results."""
    user_id = current_user.id

    # Check usage limits
    try:
        UsageService.check_limit(db, user_id, "search")
    except HTTPException as e:
        raise e

    from archive_search_service import search_developers, count_developers, developer_to_dict

    role = search.role
    location = search.location
    languages = search.languages
    min_score = search.min_score or 0
    page = search.page or 1
    per_page = search.per_page or 20
    has_email = getattr(search, 'has_email', None)
    min_experience = getattr(search, 'min_experience', None)

    # Get total count
    total = count_developers(
        db, role=role, location=location, languages=languages, min_score=min_score,
        has_email=has_email, min_experience=min_experience
    )

    # Get paginated results
    offset = (page - 1) * per_page
    results = search_developers(
        db, role=role, location=location, languages=languages,
        limit=per_page, offset=offset, min_score=min_score,
        has_email=has_email, min_experience=min_experience
    )

    # Convert to dicts
    profiles = []
    for dev in results:
        d = developer_to_dict(dev)
        # Show first name only for non-unlocked profiles
        if d.get("name"):
            d["name"] = d["name"].split()[0] if d["name"] else dev.github_username
        profiles.append(d)

    # Log usage
    UsageService.log_usage(db, user_id, "search", {"role": role, "location": location})

    return {
        "profiles": profiles,
        "total": total,
        "page": page,
        "per_page": per_page,
        "total_pages": math.ceil(total / per_page) if total > 0 else 0,
    }


# ===== GET ALL PROFILES =====

@app.get("/api/profiles", response_model=List[ProfileResponse])
def get_all_profiles(
    current_user: CurrentUser,
    db: DbSession,
    min_score: int = 0,
    max_score: int = 100,
    min_stars: int = 0,
    has_email: bool = None,
    location: str = None,
    language: str = None,
    active_within_days: int = None,
    sort_by: str = "score",
    limit: int = 100,
):
    """Get all profiles with optional filters and sorting"""
    query = db.query(Profile)
    
    # Apply filters
    if min_score > 0:
        query = query.filter(Profile.developer_score >= min_score)
    
    if max_score < 100:
        query = query.filter(Profile.developer_score <= max_score)
    
    if min_stars > 0:
        query = query.filter(Profile.total_stars >= min_stars)
    
    if has_email is not None:
        if has_email:
            query = query.filter(Profile.email.isnot(None), Profile.email != "")
        else:
            from sqlalchemy import or_
            query = query.filter(or_(Profile.email.is_(None), Profile.email == ""))
    
    if location:
        from filter_service import FilterService
        location_lower = location.lower().strip()
        location_terms = [location]
        for country, cities in FilterService.COUNTRY_CITIES.items():
            if location_lower == country or location_lower in country or country in location_lower:
                location_terms.extend(cities)
                break
        from filter_service import (
            COUNTRY_UNITED_STATES, COUNTRY_UNITED_KINGDOM, COUNTRY_UNITED_ARAB_EMIRATES
        )
        COUNTRY_ALIASES = {"usa": COUNTRY_UNITED_STATES, "us": COUNTRY_UNITED_STATES, "uk": COUNTRY_UNITED_KINGDOM, "uae": COUNTRY_UNITED_ARAB_EMIRATES}
        alias_country = COUNTRY_ALIASES.get(location_lower)
        if alias_country and alias_country in FilterService.COUNTRY_CITIES:
            location_terms.append(alias_country)
            location_terms.extend(FilterService.COUNTRY_CITIES[alias_country])
        location_terms = list(set(location_terms))
        from sqlalchemy import or_ as or_loc
        loc_filters = [Profile.location.ilike(f"%{term}%") for term in location_terms]
        query = query.filter(or_loc(*loc_filters))

    if language:
        query = query.filter(Profile.primary_language.ilike(language))
    
    if active_within_days:
        cutoff_date = datetime.now(timezone.utc) - timedelta(days=active_within_days)
        query = query.filter(Profile.last_active_date >= cutoff_date)
    
    # Apply sorting
    if sort_by == "score":
        query = query.order_by(Profile.developer_score.desc())
    elif sort_by == "stars":
        query = query.order_by(Profile.total_stars.desc())
    elif sort_by == "repos":
        query = query.order_by(Profile.public_repos.desc())
    elif sort_by == "activity":
        query = query.order_by(Profile.last_active_date.desc())
    else:
        query = query.order_by(Profile.developer_score.desc())
    
    query = query.limit(limit)
    profiles = query.all()

    return profiles


# ===== GET SINGLE PROFILE =====

@app.get("/api/profiles/{profile_id}", response_model=ProfileResponse)
def get_profile_details(
    profile_id: int,
    current_user: CurrentUser,
    db: DbSession,
):
    """Get full details for a specific profile"""
    
    profile = db.query(Profile).filter(Profile.id == profile_id).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail=f"Profile with ID {profile_id} not found")
    
    return profile


# ===== LOG PROFILE UNLOCK =====

@app.post("/api/profiles/{profile_id}/log-unlock")
def log_profile_unlock(
    profile_id: int,
    current_user: CurrentUser,
    db: DbSession,
):
    """Log that a user unlocked/viewed a profile (increments profile_views usage)"""
    user_id = current_user.id

    # Check if within limits
    UsageService.check_limit(db, user_id, "profile_view")

    # Log the unlock
    UsageService.log_usage(db, user_id, "profile_view", {"profile_id": profile_id})

    # Record unlock for smart refresh priority
    try:
        from models import ProfileUnlock
        unlock_record = ProfileUnlock(user_id=user_id, profile_id=profile_id)
        db.add(unlock_record)
        db.commit()
    except Exception:
        db.rollback()

    # Return updated usage stats
    stats = UsageService.get_usage_stats(db, user_id)
    return {
        "success": True,
        "profile_id": profile_id,
        "usage": stats.get("usage", {})
    }


# ===== TOGGLE PROFILE SELECTION =====

@app.patch("/api/profiles/{profile_id}/toggle-select")
def toggle_profile_selection(
    profile_id: int,
    current_user: CurrentUser,
    db: DbSession,
):
    """Toggle selection status of a profile"""
    
    profile = db.query(Profile).filter(Profile.id == profile_id).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail=f"Profile with ID {profile_id} not found")
    
    profile.selected = not profile.selected
    db.commit()
    db.refresh(profile)
    
    return {
        "id": profile.id,
        "username": profile.github_username,
        "selected": profile.selected
    }


# ===== GET SELECTED PROFILES =====

@app.get("/api/selected-profiles", response_model=List[ProfileResponse])
def get_selected_profiles(
    current_user: CurrentUser,
    db: DbSession,
):
    """Get all profiles marked as selected"""
    
    profiles = db.query(Profile).filter(Profile.selected == True).all()
    return profiles


# ===== SEND BULK EMAILS (UPDATED FOR RESEND) =====

def _validate_email_usage(db, user_id, profile_count):
    """Check email limits, trial expiry, and raise HTTPException if exceeded."""
    usage = UsageService.check_email_limit(db, user_id)

    # Check trial expiry
    if usage.get("trial_expired"):
        raise HTTPException(
            status_code=403,
            detail={
                "error": "TRIAL_EXPIRED",
                "message": "Your free trial has expired. Please upgrade to continue sending emails."
            }
        )

    if not usage["can_send"]:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "EMAIL_LIMIT_EXCEEDED",
                "message": f"Email limit reached. You've used {usage['used']}/{usage['limit']} emails.",
                "usage": usage
            }
        )

    if profile_count > usage["remaining"]:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "EMAIL_LIMIT_EXCEEDED",
                "message": f"Cannot send {profile_count} emails. Only {usage['remaining']} remaining.",
                "usage": usage
            }
        )


def _log_sent_emails(db, user, results, user_settings):
    """Log successfully sent emails to the EmailOutreach table."""
    for detail in results['details']:
        if detail['status'] != 'sent':
            continue
        profile = db.query(Profile).filter(Profile.id == detail['profile_id']).first()
        if not profile:
            continue
        outreach = EmailOutreach(
            user_id=user.id,
            profile_id=profile.id,
            subject=user_settings['email_subject'],
            body=user_settings['email_template'],
            status='sent',
            sent_at=datetime.now(timezone.utc)
        )
        db.add(outreach)
    db.commit()


@app.post("/api/send-bulk-emails")
async def send_bulk_emails_endpoint(
    request: dict,
    current_user: CurrentUser,
    db: DbSession,
):
    """Send bulk emails to selected profiles using Resend"""
    try:
        profile_ids = request.get("profile_ids", [])

        if not profile_ids:
            raise HTTPException(status_code=400, detail="No profiles selected")

        user_settings = {
            'sender_email': current_user.sender_email or 'noreply@talentbox.co',
            'sender_name': current_user.sender_name or current_user.name,
            'email_subject': current_user.email_subject or 'Exciting Opportunity',
            'email_template': current_user.email_template or 'Hi {{name}},\n\nWe found your profile interesting!',
            'reply_method': current_user.reply_method or 'email',
            'reply_link': current_user.reply_link or ''
        }

        _validate_email_usage(db, current_user.id, len(profile_ids))

        # Reserve quota upfront to prevent race-condition over-usage
        # (if two requests come in simultaneously, both must claim quota before sending)
        reserved_count = len(profile_ids)
        current_user.usage_emails_sent = (current_user.usage_emails_sent or 0) + reserved_count
        db.commit()

        profiles = db.query(Profile).filter(Profile.id.in_(profile_ids)).all()
        profiles_data = [{
            'id': p.id,
            'name': p.name,
            'github_username': p.github_username,
            'email': p.email
        } for p in profiles]

        results = EmailService.send_bulk_emails(profiles_data, user_settings)

        _log_sent_emails(db, current_user, results, user_settings)

        # Adjust quota: give back credits for failed emails
        if results['failed'] > 0:
            current_user.usage_emails_sent = max(0, (current_user.usage_emails_sent or 0) - results['failed'])
            db.commit()

        return {
            "success": True,
            "sent": results['sent'],
            "failed": results['failed'],
            "message": f"Successfully sent {results['sent']} emails. {results['failed']} failed."
        }

    except HTTPException:
        raise
    except Exception as e:
        # The quota reservation was already committed, so rollback() won't undo it.
        # Re-fetch the user and subtract the reserved count to restore their quota.
        try:
            db.rollback()
            fresh_user = db.query(User).filter(User.id == current_user.id).first()
            if fresh_user:
                fresh_user.usage_emails_sent = max(0, (fresh_user.usage_emails_sent or 0) - reserved_count)
                db.commit()
        except Exception as restore_err:
            logger.error(f"Failed to restore email quota after bulk send error: {restore_err}")
        logger.error(f"Bulk email error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

# ===== FILTER BY SCORE =====

@app.post("/api/filter-by-score")
def filter_by_score(
    profile_ids: List[int],
    current_user: CurrentUser,
    db: DbSession,
    min_score: int = 0,
    max_score: int = 100,
):
    """Filter profiles by developer score"""
    
    profiles = db.query(Profile).filter(Profile.id.in_(profile_ids)).all()
    filtered = FilterService.filter_by_score(profiles, min_score, max_score)
    
    return {
        "total": len(filtered),
        "profiles": filtered
    }


# ===== GET USAGE STATS =====

@app.get("/api/usage-stats")
def get_usage_stats(
    current_user: CurrentUser,
    db: DbSession,
):
    """Get current usage statistics for user"""
    user_id = current_user.id
    stats = UsageService.get_usage_stats(db, user_id)
    return stats

# ===== CHECK CSV EXPORT LIMIT =====

@app.get("/api/check-csv-limit")
def check_csv_limit(
    current_user: CurrentUser,
    db: DbSession,
):
    """Check if user can export CSV"""
    usage = UsageService.check_csv_limit(db, current_user.id)
    return usage


# ===== LOG CSV EXPORT =====

@app.post("/api/log-csv-export")
def log_csv_export(
    current_user: CurrentUser,
    db: DbSession,
):
    """Log a CSV export"""
    UsageService.log_csv_export(db, current_user.id)
    return {"success": True}

# ===== GREENHOUSE OAUTH CALLBACK (placeholder) =====

@app.get("/greenhouse/callback")
def greenhouse_callback(code: Optional[str] = None, error: Optional[str] = None):
    """Placeholder Greenhouse OAuth callback endpoint."""
    if error:
        return {"status": "error", "message": error}
    if code:
        logger.info(f"Greenhouse OAuth callback received code: {code}")
        return {"status": "success", "code": code}
    return {"status": "error", "message": "no code received"}


# ===== HEALTH CHECK =====

@app.get("/api/health")
def health_check():
    """Check if API is running"""
    return {
        "status": "healthy",
        "version": "2.0.1",
        "timestamp": datetime.now(timezone.utc).isoformat()
    }


@app.get("/api/health/deep")
def deep_health_check(db: DbSession):
    """Check if API and database are running"""
    try:
        # Test database connection
        from sqlalchemy import text
        db.execute(text("SELECT 1"))
        db_status = "healthy"
    except Exception as e:
        logger.error(f"Database health check failed: {e}")
        db_status = "unhealthy"
    
    return {
        "status": "healthy" if db_status == "healthy" else "degraded",
        "database": db_status,
        "version": "2.0.1",
        "timestamp": datetime.now(timezone.utc).isoformat()
    }


# ===== VERSION CHECK (for debugging) =====

@app.get("/api/version-check")
def version_check():
    """Verify which version of code is running - for debugging only"""
    return {
        "version": "2.0.1_RESEND",
        "features": [
            "Resend email integration",
            "Sender name support",
            "Email subject customization",
            "Reply method (email/form)",
            "UTM tracking for forms",
            "{{name}} variable with first name only"
        ],
        "fixes_applied": [
            "Fix #3: Empty profiles array bug",
            "Fix #4: Enhanced logging",
            "Fix #5: FilterService debug logs",
            "Fix #6: Profile saving error handling",
            "Fix #7: Missing logging import"
        ],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "status": "Resend Integration Complete ✅"
    }


# ===== ADMIN USAGE ENDPOINT =====

@app.get("/admin/usage")
async def admin_usage(
    request: Request,
    db: DbSession,
    days: int = 7,
    admin_key: str = None,
):
    """Top users by API consumption. Requires ADMIN_SECRET_KEY query param."""
    expected_key = os.getenv("ADMIN_SECRET_KEY", "")
    if not expected_key or admin_key != expected_key:
        raise HTTPException(status_code=403, detail="Admin access denied")

    from rate_limit_service import get_admin_usage_stats
    return get_admin_usage_stats(db, days=days)


# ===== AI SERVICE ENDPOINTS =====

@app.get("/api/ai-status")
async def ai_status_endpoint():
    """Check AI service status."""
    return check_groq_status()


@app.post("/api/parse-job-description")
async def parse_jd_endpoint(
    request: Request,
    current_user: CurrentUser,
    db: DbSession,
):
    """Parse a job description and return suggested search filters."""
    try:
        body = await request.json()
        jd_text = body.get("job_description", "")

        if not jd_text or len(jd_text) < 50:
            raise HTTPException(status_code=400, detail="Job description too short (min 50 characters)")

        if len(jd_text) > 10000:
            raise HTTPException(status_code=400, detail="Job description too long (max 10,000 characters)")

        result = await parse_job_description(jd_text)

        if not result:
            raise HTTPException(status_code=500, detail="Failed to parse job description. AI service may be unavailable.")

        return {"success": True, "filters": result}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"JD parse error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/extract-jd-file")
async def extract_jd_file(file: UploadFile = File(...)):
    """Extract plain text from an uploaded JD file (.pdf, .docx, .txt)."""
    import io
    try:
        filename = file.filename or ""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext not in ("pdf", "docx", "doc", "txt"):
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT.")

        raw = await file.read()
        if len(raw) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large (max 5 MB).")

        text = ""
        if ext == "txt":
            text = raw.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(raw))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as pdf_err:
                logger.error(f"PDF extraction error: {pdf_err}")
                raise HTTPException(status_code=422, detail="Could not extract text from PDF.")
        elif ext in ("docx", "doc"):
            try:
                import docx as _docx
                doc = _docx.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as docx_err:
                logger.error(f"DOCX extraction error: {docx_err}")
                raise HTTPException(status_code=422, detail="Could not extract text from DOCX.")

        text = text.strip()
        if len(text) < 30:
            raise HTTPException(status_code=422, detail="File appears to be empty or unreadable.")

        return {
            "success": True,
            "text": text,
            "filename": filename,
            "char_count": len(text),
            "word_count": len(text.split()),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"JD file extract error: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract file contents.")


@app.post("/api/generate-profile-summary")
async def generate_summary_endpoint(
    request: Request,
    current_user: CurrentUser,
    db: DbSession,
):
    """Generate AI summary for a single profile."""
    try:
        body = await request.json()
        profile_data = body.get("profile", {})

        if not profile_data:
            raise HTTPException(status_code=400, detail="Profile data required")

        summary = await generate_profile_summary(profile_data)

        if not summary:
            return {"success": False, "summary": None, "message": "AI service unavailable"}

        return {"success": True, "summary": summary}

    except Exception as e:
        logger.error(f"Summary generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== DEVCARD & HIRE-FREE MODELS =====

class DevCardRequest(BaseModel):
    github_username: str
    email: Optional[str] = None
    linkedin_url: Optional[str] = None
    name: Optional[str] = None
    phone: Optional[str] = None


class FreeHireRequest(BaseModel):
    company_email: str
    company_name: Optional[str] = None
    contact_name: Optional[str] = None
    job_title: str
    job_description: Optional[str] = None
    required_skills: Optional[List[str]] = None
    preferred_location: Optional[str] = None
    experience_min: Optional[int] = 0
    remote_ok: Optional[bool] = False
    jd_source: Optional[str] = 'pasted'
    jd_filename: Optional[str] = None


# ===== DEVCARD ENDPOINTS =====

def _calculate_devcard_score(followers: int, public_repos: int, total_stars: int, bio: str, email: str, primary_languages: list) -> int:
    """Replicate GithubDeveloper.calculate_score logic for DevCard profiles."""
    score = 0

    if followers:
        if followers >= 1000:
            score += 25
        elif followers >= 500:
            score += 20
        elif followers >= 100:
            score += 15
        elif followers >= 50:
            score += 10
        elif followers >= 10:
            score += 5

    if public_repos:
        if public_repos >= 50:
            score += 20
        elif public_repos >= 30:
            score += 15
        elif public_repos >= 15:
            score += 10
        elif public_repos >= 5:
            score += 5

    if total_stars:
        if total_stars >= 500:
            score += 25
        elif total_stars >= 100:
            score += 20
        elif total_stars >= 50:
            score += 15
        elif total_stars >= 10:
            score += 10
        elif total_stars >= 1:
            score += 5

    # Bio present (5 points)
    if bio and len(bio.strip()) > 20:
        score += 5

    # Email present (5 points)
    if email:
        score += 5

    # Multiple languages (5 points)
    if primary_languages and len(primary_languages) >= 3:
        score += 5

    # Recent activity assumed (max 15 points) — cap at 10 for GitHub API data
    score += 10

    return min(score, 100)


# HIDDEN: DevCard endpoints — kept for future reference
# @app.post("/api/devcard/generate")
async def generate_devcard(req: DevCardRequest, db: DbSession):
    """Generate a DevCard from a GitHub username. Public endpoint — no auth required."""
    import httpx
    import re

    try:
        # Clean username: strip @, extract from URL
        username = req.github_username.strip()
        username = re.sub(r'^@', '', username)
        url_match = re.search(r'github\.com/([^/\s]+)', username)
        if url_match:
            username = url_match.group(1)
        username = username.strip('/')

        if not username or not re.match(r'^[a-zA-Z0-9]([a-zA-Z0-9\-]{0,37}[a-zA-Z0-9])?$', username):
            raise HTTPException(status_code=400, detail="Invalid GitHub username.")

        from sqlalchemy import text as sql_text

        # Check cache (updated within last 24 hours)
        cache_row = db.execute(sql_text("""
            SELECT * FROM devcard_profiles
            WHERE github_username = :username
              AND updated_at >= NOW() - INTERVAL '24 hours'
              AND is_published = TRUE
        """), {"username": username}).fetchone()

        if cache_row:
            db.execute(sql_text(
                "UPDATE devcard_profiles SET views_count = views_count + 1 WHERE github_username = :username"
            ), {"username": username})
            db.commit()
            row = dict(cache_row._mapping)
            return {"success": True, "card": row}

        # Fetch from GitHub API
        gh_headers = {
            "Authorization": f"token {GITHUB_TOKEN}",
            "Accept": "application/vnd.github.v3+json",
        }
        async with httpx.AsyncClient() as client:
            user_resp = await client.get(
                f"https://api.github.com/users/{username}",
                headers=gh_headers, timeout=10.0
            )
            if user_resp.status_code == 404:
                raise HTTPException(status_code=404, detail="GitHub user not found.")
            if user_resp.status_code != 200:
                raise HTTPException(status_code=502, detail="GitHub API unavailable. Please try again.")
            user_data = user_resp.json()

            if user_data.get("type") == "Organization":
                raise HTTPException(status_code=400, detail="Organization accounts are not supported.")

            repos_resp = await client.get(
                f"https://api.github.com/users/{username}/repos",
                headers=gh_headers,
                params={"sort": "stars", "per_page": 10},
                timeout=10.0
            )
            repos = repos_resp.json() if repos_resp.status_code == 200 else []

        # Language percentages from repos
        lang_bytes: dict = {}
        for repo in repos:
            lang = repo.get("language")
            if lang:
                lang_bytes[lang] = lang_bytes.get(lang, 0) + (repo.get("size", 1) or 1)

        total_bytes = sum(lang_bytes.values()) or 1
        sorted_langs = sorted(lang_bytes.items(), key=lambda x: x[1], reverse=True)[:7]
        language_percentages = {lang: round(b / total_bytes * 100, 1) for lang, b in sorted_langs}
        primary_languages = list(language_percentages.keys())

        # Top 3 projects by stars
        top_projects = []
        for repo in sorted(repos, key=lambda r: r.get("stargazers_count", 0), reverse=True)[:3]:
            desc = repo.get("description") or ""
            top_projects.append({
                "name": repo.get("name"),
                "description": desc[:80] if desc else None,
                "stars": repo.get("stargazers_count", 0),
                "language": repo.get("language"),
            })

        # Contribution stats
        total_stars = sum(r.get("stargazers_count", 0) for r in repos)
        public_repos = user_data.get("public_repos", 0)
        followers = user_data.get("followers", 0)
        created_at_str = user_data.get("created_at", "")
        try:
            account_year = int(created_at_str[:4])
            years_active = datetime.now().year - account_year
        except Exception:
            years_active = 1
        estimated_experience_years = max(1, years_active)

        contribution_stats = {
            "stars": total_stars,
            "repos": public_repos,
            "followers": followers,
            "years_active": years_active,
        }

        # AI analysis via GPT-4o-mini
        ai_summary = None
        detected_role = "Software Developer"
        seniority_level = "Mid-Level"

        try:
            if openai.api_key:
                lang_str = ", ".join(primary_languages[:5]) if primary_languages else "unknown"
                ai_prompt = f"""Analyze this GitHub developer profile and return JSON only (no markdown):
Username: {username}
Name: {user_data.get('name', username)}
Bio: {user_data.get('bio', 'N/A')}
Location: {user_data.get('location', 'N/A')}
Languages: {lang_str}
Public Repos: {public_repos}
Followers: {followers}
Stars: {total_stars}
Years Active: {years_active}

Return JSON with exactly these fields:
{{"summary": "2-sentence recruiter brief", "role": "one of: Frontend Developer, Backend Developer, Full-Stack Developer, Mobile Developer, DevOps Engineer, Data Scientist, AI/ML Engineer, Data Engineer, Security Engineer, Software Developer", "seniority": "one of: Junior, Mid-Level, Senior, Expert"}}"""

                client_ai = openai.OpenAI(api_key=openai.api_key)
                ai_resp = client_ai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": ai_prompt}],
                    max_tokens=200,
                    temperature=0.3,
                )
                ai_text = ai_resp.choices[0].message.content.strip()
                # Strip markdown code fences if present
                ai_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', ai_text, flags=re.MULTILINE).strip()
                ai_json = json.loads(ai_text)
                ai_summary = ai_json.get("summary", "")
                detected_role = ai_json.get("role", "Software Developer")
                seniority_level = ai_json.get("seniority", "Mid-Level")
        except Exception as ai_err:
            logger.warning(f"DevCard AI analysis failed for {username}: {ai_err}")
            # Fallback summary
            lang_str = ", ".join(primary_languages[:3]) if primary_languages else "various languages"
            ai_summary = (
                f"{username} is a developer with {years_active}+ years of GitHub activity, "
                f"specialising in {lang_str}. They have {public_repos} public repositories and {total_stars} total stars."
            )
            # Fallback role detection from languages
            if "Swift" in primary_languages or "Kotlin" in primary_languages:
                detected_role = "Mobile Developer"
            elif "Python" in primary_languages and ("TensorFlow" in str(user_data.get("bio", "")) or "ML" in str(user_data.get("bio", ""))):
                detected_role = "AI/ML Engineer"
            elif any(l in primary_languages for l in ["JavaScript", "TypeScript", "HTML", "CSS"]):
                detected_role = "Frontend Developer"
            elif any(l in primary_languages for l in ["Python", "Go", "Java", "Ruby", "PHP", "C#", "Rust"]):
                detected_role = "Backend Developer"

            if years_active <= 2:
                seniority_level = "Junior"
            elif years_active <= 5:
                seniority_level = "Mid-Level"
            elif years_active <= 9:
                seniority_level = "Senior"
            else:
                seniority_level = "Expert"

        developer_score = _calculate_devcard_score(
            followers=followers,
            public_repos=public_repos,
            total_stars=total_stars,
            bio=user_data.get("bio", ""),
            email=user_data.get("email", ""),
            primary_languages=primary_languages,
        )

        display_name = user_data.get("name") or username
        avatar_url = user_data.get("avatar_url", "")
        bio = user_data.get("bio", "")
        location = user_data.get("location", "")

        # Upsert into devcard_profiles
        db.execute(sql_text("""
            INSERT INTO devcard_profiles (
                github_username, display_name, avatar_url, bio, location,
                detected_role, seniority_level, primary_languages, language_percentages,
                top_projects, contribution_stats, ai_summary,
                developer_score, estimated_experience_years,
                email, linkedin_url, phone,
                views_count, is_published, updated_at
            ) VALUES (
                :username, :display_name, :avatar_url, :bio, :location,
                :detected_role, :seniority_level, :primary_languages, :language_percentages,
                :top_projects, :contribution_stats, :ai_summary,
                :developer_score, :estimated_experience_years,
                :email, :linkedin_url, :phone,
                1, TRUE, NOW()
            )
            ON CONFLICT (github_username) DO UPDATE SET
                display_name = EXCLUDED.display_name,
                avatar_url = EXCLUDED.avatar_url,
                bio = EXCLUDED.bio,
                location = EXCLUDED.location,
                detected_role = EXCLUDED.detected_role,
                seniority_level = EXCLUDED.seniority_level,
                primary_languages = EXCLUDED.primary_languages,
                language_percentages = EXCLUDED.language_percentages,
                top_projects = EXCLUDED.top_projects,
                contribution_stats = EXCLUDED.contribution_stats,
                ai_summary = EXCLUDED.ai_summary,
                developer_score = EXCLUDED.developer_score,
                estimated_experience_years = EXCLUDED.estimated_experience_years,
                email = COALESCE(EXCLUDED.email, devcard_profiles.email),
                linkedin_url = COALESCE(EXCLUDED.linkedin_url, devcard_profiles.linkedin_url),
                phone = COALESCE(EXCLUDED.phone, devcard_profiles.phone),
                views_count = devcard_profiles.views_count + 1,
                updated_at = NOW()
        """), {
            "username": username,
            "display_name": display_name,
            "avatar_url": avatar_url,
            "bio": bio,
            "location": location,
            "detected_role": detected_role,
            "seniority_level": seniority_level,
            "primary_languages": primary_languages,
            "language_percentages": json.dumps(language_percentages),
            "top_projects": json.dumps(top_projects),
            "contribution_stats": json.dumps(contribution_stats),
            "ai_summary": ai_summary,
            "developer_score": developer_score,
            "estimated_experience_years": estimated_experience_years,
            "email": req.email or user_data.get("email"),
            "linkedin_url": req.linkedin_url,
            "phone": req.phone,
        })

        # Insert into candidate_waitlist if email was provided
        if req.email:
            try:
                db.execute(sql_text("""
                    INSERT INTO candidate_waitlist (email, name, github_username, linkedin_url, phone, source)
                    VALUES (:email, :name, :github_username, :linkedin_url, :phone, 'devcard')
                """), {
                    "email": req.email,
                    "name": req.name or display_name,
                    "github_username": username,
                    "linkedin_url": req.linkedin_url,
                    "phone": req.phone,
                })
            except Exception as wl_err:
                logger.warning(f"Waitlist insert failed for {username}: {wl_err}")

        db.commit()

        card = {
            "github_username": username,
            "display_name": display_name,
            "avatar_url": avatar_url,
            "bio": bio,
            "location": location,
            "detected_role": detected_role,
            "seniority_level": seniority_level,
            "primary_languages": primary_languages,
            "language_percentages": language_percentages,
            "top_projects": top_projects,
            "contribution_stats": contribution_stats,
            "ai_summary": ai_summary,
            "developer_score": developer_score,
            "estimated_experience_years": estimated_experience_years,
        }
        return {"success": True, "card": card}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DevCard generation error for {req.github_username}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate DevCard. Please try again.")


# HIDDEN: DevCard endpoints — kept for future reference
# @app.get("/api/devcard/{username}")
async def get_devcard(username: str, db: DbSession):
    """Fetch a published DevCard by GitHub username. Public endpoint — no auth required."""
    try:
        from sqlalchemy import text as sql_text

        row = db.execute(sql_text("""
            SELECT * FROM devcard_profiles
            WHERE github_username = :username AND is_published = TRUE
        """), {"username": username}).fetchone()

        if not row:
            raise HTTPException(status_code=404, detail="DevCard not found.")

        db.execute(sql_text(
            "UPDATE devcard_profiles SET views_count = views_count + 1 WHERE github_username = :username"
        ), {"username": username})
        db.commit()

        card = dict(row._mapping)

        # Parse JSONB fields that may come back as strings
        for field in ("language_percentages", "top_projects", "contribution_stats", "experience_history"):
            val = card.get(field)
            if isinstance(val, str):
                try:
                    card[field] = json.loads(val)
                except Exception:
                    pass

        return {"success": True, "card": card}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DevCard fetch error for {username}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to fetch DevCard.")


# ===== HIRE FREE ENDPOINT =====

def _detect_role_from_title(job_title: str) -> str:
    """Map job title keywords to a standard role string."""
    title_lower = job_title.lower()
    if "frontend" in title_lower or "front-end" in title_lower or "front end" in title_lower:
        return "Frontend Developer"
    if "backend" in title_lower or "back-end" in title_lower or "back end" in title_lower:
        return "Backend Developer"
    if "full stack" in title_lower or "fullstack" in title_lower or "full-stack" in title_lower:
        return "Full-Stack Developer"
    if "mobile" in title_lower or "ios" in title_lower or "android" in title_lower or "flutter" in title_lower:
        return "Mobile Developer"
    if "devops" in title_lower or "sre" in title_lower or "platform engineer" in title_lower:
        return "DevOps Engineer"
    if "data scientist" in title_lower:
        return "Data Scientist"
    if "machine learning" in title_lower or " ml " in title_lower or "mlops" in title_lower or "ai engineer" in title_lower:
        return "AI/ML Engineer"
    if "data engineer" in title_lower:
        return "Data Engineer"
    if "security" in title_lower:
        return "Security Engineer"
    if "qa" in title_lower or "quality" in title_lower or "test engineer" in title_lower:
        return "QA Engineer"
    if "blockchain" in title_lower or "web3" in title_lower or "solidity" in title_lower:
        return "Blockchain Developer"
    if "embedded" in title_lower or "firmware" in title_lower or "iot" in title_lower:
        return "Embedded Engineer"
    return "Software Developer"


@app.post("/api/hire-free")
async def hire_free(req: FreeHireRequest, db: DbSession):
    """Accept a job request, match developer profiles, and email results. Public endpoint."""
    import httpx

    try:
        from sqlalchemy import text as sql_text
        from archive_search_service import search_developers, developer_to_dict

        import hashlib as _hashlib, re as _re

        # 1. Insert hire request with status='processing'
        jd_text = req.job_description or ""
        result = db.execute(sql_text("""
            INSERT INTO free_hire_requests
                (company_email, company_name, contact_name, job_title, job_description,
                 required_skills, preferred_location, experience_min, remote_ok, status,
                 jd_source, jd_filename, jd_word_count, jd_char_count)
            VALUES
                (:company_email, :company_name, :contact_name, :job_title, :job_description,
                 :required_skills, :preferred_location, :experience_min, :remote_ok, 'processing',
                 :jd_source, :jd_filename, :jd_word_count, :jd_char_count)
            RETURNING id
        """), {
            "company_email": req.company_email,
            "company_name": req.company_name,
            "contact_name": req.contact_name,
            "job_title": req.job_title,
            "job_description": jd_text or None,
            "required_skills": req.required_skills or [],
            "preferred_location": req.preferred_location,
            "experience_min": req.experience_min or 0,
            "remote_ok": req.remote_ok or False,
            "jd_source": req.jd_source or "pasted",
            "jd_filename": req.jd_filename,
            "jd_word_count": len(jd_text.split()) if jd_text else None,
            "jd_char_count": len(jd_text) if jd_text else None,
        })
        db.commit()
        request_id = result.fetchone()[0]

        # 2. Detect role from job title
        detected_role = _detect_role_from_title(req.job_title)

        # 3. Extract skills + parse JD via AI if provided (with Redis cache)
        all_skills = list(req.required_skills or [])
        parsed_role = None
        parsed_location = None
        parsed_experience = None

        if jd_text and openai.api_key:
            jd_hash = _hashlib.md5(jd_text[:5000].encode()).hexdigest()
            cached_parse = await get_cached_jd_parse(jd_hash)

            if cached_parse:
                logger.info(f"JD parse cache hit: {jd_hash}")
                ai_skills = cached_parse.get("skills", [])
                parsed_role = cached_parse.get("role")
                parsed_location = cached_parse.get("location")
                parsed_experience = cached_parse.get("experience")
            else:
                try:
                    client_ai = openai.OpenAI(api_key=openai.api_key)
                    skill_resp = client_ai.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": (
                            "Analyze this job description and return a JSON object with:\n"
                            "- skills: array of top 5-8 technical skills/technologies\n"
                            "- role: primary developer role (one of: Frontend Developer, Backend Developer, "
                            "Full-Stack Developer, Mobile Developer, DevOps Engineer, Data Scientist, "
                            "AI/ML Engineer, Data Engineer, Security Engineer, QA Engineer, "
                            "Blockchain Developer, Game Developer, Embedded Engineer, Software Developer)\n"
                            "- location: preferred location string or null\n"
                            "- experience: minimum years of experience as integer or null\n"
                            "Return ONLY valid JSON, no markdown.\n\n"
                            f"{jd_text[:3000]}"
                        )}],
                        max_tokens=250,
                        temperature=0.1,
                    )
                    raw_text = skill_resp.choices[0].message.content.strip()
                    raw_text = _re.sub(r'^```(?:json)?\s*|\s*```$', '', raw_text, flags=_re.MULTILINE).strip()
                    parsed = json.loads(raw_text)
                    ai_skills = parsed.get("skills", []) if isinstance(parsed.get("skills"), list) else []
                    parsed_role = parsed.get("role")
                    parsed_location = parsed.get("location")
                    parsed_experience = parsed.get("experience")
                    await set_cached_jd_parse(jd_hash, {
                        "skills": ai_skills,
                        "role": parsed_role,
                        "location": parsed_location,
                        "experience": parsed_experience,
                    })
                except Exception as skill_err:
                    logger.warning(f"AI JD parse failed: {skill_err}")
                    ai_skills = []

            for s in ai_skills:
                if isinstance(s, str) and s not in all_skills:
                    all_skills.append(s)

            # Persist parsed fields to DB
            if any([parsed_role, parsed_location, parsed_experience, ai_skills]):
                try:
                    db.execute(sql_text("""
                        UPDATE free_hire_requests SET
                            jd_parsed_role = :role,
                            jd_parsed_skills = :skills,
                            jd_parsed_location = :location,
                            jd_parsed_experience = :experience
                        WHERE id = :id
                    """), {
                        "role": parsed_role,
                        "skills": ai_skills if ai_skills else None,
                        "location": parsed_location,
                        "experience": parsed_experience,
                        "id": request_id,
                    })
                    db.commit()
                except Exception as db_err:
                    logger.warning(f"Failed to persist parsed JD fields: {db_err}")

        # 4. Search DB for matching developers (use AI-parsed values when available)
        search_role = parsed_role or detected_role
        search_location = req.preferred_location or parsed_location
        search_experience = req.experience_min or parsed_experience
        db_results = search_developers(
            db,
            role=search_role,
            location=search_location,
            languages=all_skills if all_skills else None,
            limit=20,
            min_score=30,
            has_email=True,
            min_experience=search_experience if search_experience else None,
        )
        matched_dicts = [developer_to_dict(dev) for dev in db_results]

        # 5. Supplement with live GitHub API if fewer than 15 results
        if len(matched_dicts) < 15:
            gh_headers = {
                "Authorization": f"token {GITHUB_TOKEN}",
                "Accept": "application/vnd.github.v3+json",
            }
            primary_skill = all_skills[0] if all_skills else ""
            search_location = req.preferred_location or ""

            queries_to_try = []
            if primary_skill and search_location:
                queries_to_try.append(f"location:{search_location} language:{primary_skill} repos:>3")
            if primary_skill:
                queries_to_try.append(f"language:{primary_skill} repos:>5 followers:>10")
            if search_location:
                queries_to_try.append(f"location:{search_location} repos:>5")

            seen_usernames = {p.get("github_username") for p in matched_dicts}

            async with httpx.AsyncClient() as client:
                for query in queries_to_try:
                    if len(matched_dicts) >= 15:
                        break
                    try:
                        resp = await client.get(
                            "https://api.github.com/search/users",
                            headers=gh_headers,
                            params={"q": query, "per_page": 20, "sort": "followers"},
                            timeout=10.0,
                        )
                        if resp.status_code == 200:
                            users = resp.json().get("items", [])
                            for user in users:
                                if len(matched_dicts) >= 15:
                                    break
                                login = user.get("login")
                                if login in seen_usernames or user.get("type") == "Organization":
                                    continue
                                seen_usernames.add(login)
                                try:
                                    detail_resp = await client.get(
                                        f"https://api.github.com/users/{login}",
                                        headers=gh_headers, timeout=8.0
                                    )
                                    if detail_resp.status_code == 200:
                                        detail = detail_resp.json()
                                        try:
                                            acct_year = int(detail.get("created_at", "2020")[:4])
                                            exp_years = max(1, datetime.now().year - acct_year)
                                        except Exception:
                                            exp_years = 1
                                        matched_dicts.append({
                                            "id": 0,
                                            "github_username": detail.get("login"),
                                            "name": detail.get("name") or detail.get("login"),
                                            "bio": detail.get("bio"),
                                            "email": detail.get("email"),
                                            "avatar_url": detail.get("avatar_url"),
                                            "profile_url": detail.get("html_url"),
                                            "location": detail.get("location", ""),
                                            "detected_role": detected_role,
                                            "detected_roles": [],
                                            "languages": [],
                                            "public_repos": detail.get("public_repos", 0),
                                            "followers": detail.get("followers", 0),
                                            "total_stars": 0,
                                            "developer_score": 50,
                                            "estimated_experience_years": exp_years,
                                        })
                                except Exception:
                                    pass
                    except Exception as gh_err:
                        logger.error(f"GitHub API supplement failed: {gh_err}")
                        continue

        # 6. Build preview (first 5, name blurred)
        preview_profiles = []
        for p in matched_dicts[:5]:
            full_name = p.get("name") or p.get("github_username") or "Developer"
            name_preview = (full_name[0] + "***") if full_name else "D***"
            preview_profiles.append({
                "id": p.get("id", 0),
                "detected_role": p.get("detected_role", detected_role),
                "location": p.get("location", ""),
                "developer_score": p.get("developer_score", 50),
                "languages": (p.get("languages") or [])[:3],
                "name_preview": name_preview,
            })

        # 7. Build HTML email
        matched_count = len(matched_dicts)
        profile_rows_html = ""
        for i, p in enumerate(matched_dicts[:20], start=1):
            full_name = p.get("name") or p.get("github_username") or "Developer"
            username_val = p.get("github_username", "")
            role_val = p.get("detected_role") or detected_role
            location_val = p.get("location") or "Remote"
            score_val = p.get("developer_score", 50)
            langs = (p.get("languages") or [])[:3]
            lang_str = ", ".join(langs) if langs else "Various"
            exp_val = p.get("estimated_experience_years", 1)
            stars_val = p.get("total_stars", 0)
            repos_val = p.get("public_repos", 0)
            gh_url = p.get("profile_url") or f"https://github.com/{username_val}"

            score_color = "#16a34a" if score_val >= 70 else ("#ea580c" if score_val >= 50 else "#6b7280")
            profile_rows_html += f"""
            <div style="background:#f9fafb;border-radius:10px;padding:16px 20px;margin-bottom:12px;border:1px solid #e5e7eb;">
                <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                    <div>
                        <span style="font-size:12px;color:#9ca3af;font-weight:600;">#{i}</span>
                        <span style="font-size:15px;font-weight:700;color:#111827;margin-left:8px;">{full_name}</span>
                        <span style="font-size:12px;color:#6b7280;margin-left:8px;">@{username_val}</span>
                    </div>
                    <span style="background:{score_color};color:#fff;font-size:12px;font-weight:700;padding:3px 10px;border-radius:100px;">Score {score_val}</span>
                </div>
                <div style="margin-top:8px;font-size:13px;color:#374151;">
                    <strong>{role_val}</strong> &bull; {location_val}
                </div>
                <div style="margin-top:6px;font-size:12px;color:#6b7280;">
                    Languages: {lang_str} &bull; {exp_val}+ yrs &bull; {stars_val} stars &bull; {repos_val} repos
                </div>
                <div style="margin-top:10px;">
                    <a href="{gh_url}" style="display:inline-block;background:#FF6B35;color:#fff;font-size:12px;font-weight:600;padding:6px 14px;border-radius:8px;text-decoration:none;">View Profile →</a>
                </div>
            </div>"""

        contact_name_str = req.contact_name or "there"
        email_html = f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"></head>
        <body style="margin:0;padding:0;font-family:'Segoe UI',Arial,sans-serif;background:#f3f4f6;">
            <div style="max-width:640px;margin:32px auto;background:#fff;border-radius:16px;overflow:hidden;box-shadow:0 4px 24px rgba(0,0,0,0.08);">
                <!-- Header -->
                <div style="background:linear-gradient(135deg,#FF6B35,#e85d27);padding:32px 36px;">
                    <div style="font-size:22px;font-weight:800;color:#fff;letter-spacing:-0.5px;">TalentBox</div>
                    <div style="font-size:16px;color:rgba(255,255,255,0.9);margin-top:8px;font-weight:500;">Your matched developer profiles are ready</div>
                </div>
                <!-- Body -->
                <div style="padding:28px 36px;">
                    <p style="font-size:15px;color:#374151;margin:0 0 6px;">Hi {contact_name_str},</p>
                    <p style="font-size:14px;color:#6b7280;margin:0 0 20px;">
                        We found <strong style="color:#111827;">{matched_count} developer{'' if matched_count == 1 else 's'}</strong> matching your role: <strong style="color:#FF6B35;">{req.job_title}</strong>
                    </p>
                    {profile_rows_html}
                </div>
                <!-- Footer CTA -->
                <div style="background:#f9fafb;border-top:1px solid #e5e7eb;padding:24px 36px;text-align:center;">
                    <p style="font-size:13px;color:#6b7280;margin:0 0 14px;">
                        Want to search <strong>200,000+</strong> developers with advanced filters?
                    </p>
                    <a href="https://talentbox.co/signup" style="display:inline-block;background:#FF6B35;color:#fff;font-weight:700;font-size:14px;padding:12px 28px;border-radius:10px;text-decoration:none;">
                        Start sourcing free →
                    </a>
                    <p style="font-size:11px;color:#9ca3af;margin:16px 0 0;">
                        TalentBox &bull; Developer Sourcing Platform
                    </p>
                </div>
            </div>
        </body>
        </html>"""

        # 8. Send email
        email_subject = f"TalentBox: {matched_count} developer{'s' if matched_count != 1 else ''} matched for '{req.job_title}'"
        try:
            EmailService.send_single_email(
                to_email=req.company_email,
                subject=email_subject,
                body_html=email_html,
            )
            email_sent = True
        except Exception as email_err:
            logger.error(f"hire-free email send failed: {email_err}")
            email_sent = False

        # 9. Update hire request record
        db.execute(sql_text("""
            UPDATE free_hire_requests SET
                matched_profiles_count = :count,
                email_sent = :email_sent,
                email_sent_at = CASE WHEN :email_sent THEN NOW() ELSE NULL END,
                status = :status
            WHERE id = :id
        """), {
            "count": matched_count,
            "email_sent": email_sent,
            "status": "sent" if email_sent else "matched",
            "id": request_id,
        })
        db.commit()

        return {
            "success": True,
            "request_id": request_id,
            "matched_count": matched_count,
            "preview_profiles": preview_profiles,
            "message": (
                f"We found {matched_count} matching developers! Check your inbox at {req.company_email}."
                if email_sent
                else f"We found {matched_count} matching developers. Email delivery is pending."
            ),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"hire-free error: {e}", exc_info=True)
        try:
            db.rollback()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Failed to process hire request. Please try again.")


# ===== RUN WITH UVICORN =====

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
